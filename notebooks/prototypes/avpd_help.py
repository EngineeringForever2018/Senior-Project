#Hayden Donovan
#University of Nevada-Reno
#CS426 Team 25 (Gage Christensen, Daniel Enriquez, Noah Wong, Jared Lam)
#AVPD
#4/11/2021

import spacy
import numpy as np
from numpy import linalg as LA
import pickle
import os
import docx
from docx.enum.text import WD_COLOR_INDEX
import PyPDF2
import sys

nlp = spacy.load("en_core_web_sm")

posVector = ['ADJ', 'ADP', 'ADV', 'AUX', 'CCONJ', 'DET', 'INTJ', 'NOUN', 'NUM', 'PART', 'PRON', 'PROPN', 'PUNCT', 'SCONJ', 'SPACE', 'SYM', 'VERB', 'X']
Phi_size = len(posVector)*len(posVector)


#TERM was removed from posVector because well... it's always going to end with a punctuation that's kind of a given
#ideally, i want to allow users to submit entire paragraphs for this to work, so it'll work on a paragraph basis
#but also let it work on smaller scale defined by us later on which means i need:

#say which sentences to hl beforehand:
    #get a list sentences
    #have an adjacent list of booleans whether or not to highlight
def decoupledMakeHLDocx(bool_list, string_list, dest_name):
    doc = docx.Document()
    para = doc.add_paragraph('')
    for i in range(len(string_list)):
        if i < len(bool_list) and bool_list[i] == True:
            para.add_run(string_list[i] + ' ').font.highlight_color = WD_COLOR_INDEX.GRAY_25
        else:
            para.add_run(string_list[i] + ' ')

    doc.save(dest_name)

    return True


def makeHighLightDocx(HL_treshold, inputStringArr, xBar, PhiHat):
    doc = docx.Document()
    para = doc.add_paragraph('')
    for i in inputStringArr:
        if test_String(i, xBar, PhiHat) > HL_treshold:
            para.add_run(i).font.highlight_color = WD_COLOR_INDEX.RED
        else:
            para.add_run(i)
    return doc


#THIS IS THE MOST APPROPRIATE STRING DUMP -> SENTENCE PARSER USE THIS ONE
def fileToSentenceList(paragraph):
    testList=[]
    #Basic parse from first character to first punctuation
    index = 0
    if(paragraph[len(paragraph)-1]=='.' or paragraph[len(paragraph)-1]=='?' or paragraph[len(paragraph)-1]=='!')==False:
        paragraph += '.'
        #last is a -> False or False or False   ->
        #last is . -> True or False or False    ->

    while(len(paragraph) > 0 and index != -1):
        for i in range(len(paragraph)):
            if paragraph[i] > '!' and paragraph[i] < '~':
                index=i
                break


        _, index = getNextPunct(paragraph)
        sample = paragraph[:index+1]
        paragraph = paragraph[index+2:]

        doc = nlp(str(sample))
        while(len(doc) > 0 and doc[len(doc)-1].pos_ != 'PUNCT' and index != -1):
            temp, index = getNextPunct(paragraph)
            if index == -1:
                print(sample)
                break
            paragraph = paragraph[index+2:]
            sample += temp
            doc = nlp(str(sample))

        testList.append(sample)

    return testList

def getNextPunct(string):
    for i in range(len(string)):
        if string[i] == '.' or string[i] == '!' or string=='?':
            return string[:i+1], i
    return string, -1

def getSentenceList(paragraph):
    testList = []
    doc = nlp(str(paragraph))
    testList.append('')
    tl_index=0
    for i in doc:
        testList[tl_index] += i.text + ' '
        if i.pos_ == 'PUNCT' and (i.text=='.' or i.text=='!' or i.text=='?'):
            tl_index+=1
            testList.append('')

    if testList[tl_index]=='':
        testList.pop()

    return testList

def makeX_fromParagraph_2(paragraph):
    xList = []
    doc = nlp(str(paragraph))


def makeX_fromParagraph(paragraph):
    posList = [[]]
    posIndex = 0
    xList = []
    doc = nlp(str(paragraph))
    xIndex=0

    for i in range(len(doc)):
        posList[posIndex].append(doc[i].pos_)
        if i < len(doc)-1 and doc[i].pos_ == 'PUNCT' and (doc[i].text=='.' or doc[i].text=='?' or doc[i].text=='!'):
            posList.append([])
            posIndex+=1
    for i in range(len(posList)):
        xList.append(np.zeros(Phi_size))

    for i in range(len(posList)):
        for j in range(len(posList)):
            if j < len(posList[i])-1:
                bigramIndex = len(posVector)*posVector.index(str(posList[i][j])) + posVector.index(str(posList[i][j+1]))
                xList[i][bigramIndex] += 1

    #for i in range(len(posList)):
    #    if i < len(posList)-1 and (posList[i]=='PUNCT' and (doc[i].text=='.' or doc[i].text=='?' or doc[i].text=='!')==False):
    #        xList[xIndex][posVector.index(str(posList[i]))*len(posVector)+posVector.index(str(posList[i+1]))]+=1
    #    elif i < len(posList)-1 and posList[i]=='PUNCT' and (doc[i].text=='.' or doc[i].text=='?' or doc[i].text=='!'):
    #        xList.append(np.zeros(Phi_size))
    #        xIndex+=1

    return xList


class tSentence:    #for transcribed sentence. May be used for general purpose isntead of just sentences to avoid issues w/ proper noun punctuation and whatnot

    data = []
    size = 0

    def __init__(self, string):
        doc = nlp(str(string))
        for i in doc:
            self.data.append(tWord(i.text, i.pos_))
            self.size += 1

    def printPOS(self):
        for i in range(self.size):
            print(self.data[i].pos, end=' ')
        print("done")
        return True

class tWord:
    
    def __init__(self, word):
        self.word = word
        self.pos = str(nlp(word)[0].pos)

    def __init__(self, word, part):
        self.word = str(word)
        self.pos = str(part)


    def __repr__(self):
        return str('<'+ str(self.word)+", "+ str(self.pos)+'>')

    def __str__(self):
        return str('<'+ str(self.word)+", "+ str(self.pos)+'>')

#Finds ||x|| of a vector
def getNorm(vector):
    sum=0.0
    for i in vector:
        sum+=i*i
    return np.sqrt(sum)

#States what bigram based on a given index. Mostly used for debugging purposes, has little to no implementation use.
def getBigram(index):
    first = index//len(posVector)
    second = index%len(posVector)
    print(posVector[index//len(posVector)], end=' ')
    if(posVector[index%len(posVector)]=='TERM'):
        print('***', posVector[index%len(posVector)], '***')
    else:
        print(posVector[index%len(posVector)])
    return True


#Pass string, return bigram vector X (name in function is wrong)
def make_X(string):
    posArr = tSentence(string)
    Phi = np.zeros(Phi_size)

    k=0

    for i in range(posArr.size-1):
        if i == 0:
            l = posVector.index(str(posArr.data[i].pos))
            k = posVector.index(str(posArr.data[i+1].pos))
            Phi[l*len(posVector)+k] += 1
        else:
            l = posVector.index(str(posArr.data[i+1].pos))
            Phi[k*len(posVector)+l] += 1
            k=l
    return Phi

#takes list of X and compiles it into an average, xBar (the name is wrong)
def make_xBar(X_list):
    xBar = np.zeros(Phi_size)
    for i in X_list:
        xBar = xBar + i
    xBar = xBar/len(X_list)

    return xBar

#Makes a list of Phi, imagine it's stored vertically
def make_PhiList(X, xBar):
    PhiList = []
    for i in X:
        PhiList.append(i-xBar)
    return PhiList

#Want a better place to hold them while still allowing A to exist so:
def storePhi(Phi_list, Phi):
    #where Phi list is the array of Phi unformatted. A is too formatted to use cleanly at this point
    Phi_list.append(Phi)
    return True

#but i also kind of need A for actual computations
def make_A(Phi_list):
    A = np.zeros((Phi_size, len(Phi_list)))
    for i in range(len(Phi_list)):
        for j in range(len(Phi_list[i])):
            A[j][i] = Phi_list[i][j]
    return A

#AAT is the default = [324xL]x[Lx324] = [324x324]
#ATA would          = [Lx324]x[324xL] = [LxL]
def make_covarianceMatrix(A):
    SigmaX = np.matmul(A, A.transpose())
    SigmaX = SigmaX/A[0].size
    return SigmaX

#Returns eigenvector and eigenvector coefficients (v and lambda)
def make_Eigen(cov):
    lamb, v = LA.eig(cov)#assuming cov is AAT covariance matrix
    v = v.real
    lamb = lamb.real
    return lamb, v

#Takes V created by make_Eigen() and stores it in a list for easier manipulation later.
def store_v(V):
    #i believe v is stored similarly to A and Phi, so vertically
    vList = []
    for i in range(len(V[0])):
        temp = np.zeros(len(V[0]))
        for j in range(len(V)):
            temp[j] = float(V[j][i])
        vList.append(temp)

    return vList

#Normalizes V such that ||V_i||==1, allows us to easily calculate Omega
def normalize_vList(vList):
    for i in vList:
        sum=0.0
        for j in range(len(i)):
            sum += i[j]*i[j]
        sum = np.sqrt(sum)
        #print(sum)
        
        for j in range(len(i)):
            i[j] /= sum

    return vList

#Creates omega using V and Phi, used for calculating Phi Hat down the line
def make_omegaList(Phi_list, vList):
    Omega_list = []
    for i in range(len(Phi_list)):
        Omega_list.append(np.matmul(Phi_list[i].transpose(), vList[i]))
    return Omega_list

#Returns how many eigenvectors are to be used to represent at least T% of the data
def get_kThresh(lamb, T):
    denom = np.sum(lamb)
    k = 0
    sum = 0.0
    for k in range(len(lamb)):
        sum += lamb[k]
        if sum/denom > T:
            return k+1
    return k+1

#Creates sample of what the data should look like based on k samples holding T% of the data.
def get_PhiHat(k, omegaList, vList):
    PhiHat = np.zeros(len(vList[0]))
    for i in range(k):
        PhiHat += omegaList[i]*vList[i]
    return PhiHat

#Returns the euclidean distance of a given string from the current model.
def test_String(string, mean, PhiHat):
    test = make_X(string)
    test = test-mean
    #Makes an i-mean = xbar

    return getNorm(test - PhiHat)

#Takes the string and vital parts and updates everything with the new string added to the data set. Returns all passed values except string
def integrate_String(string, xList, xBar, Phi_list, A, lamb, vList, omegaList):
    newX = make_X(string)
    xList.append(newX)
    xBar = make_xBar(xList)

    Phi_list.append(newX-xBar)
    A = make_A(Phi_list)

    lamb, v = make_Eigen(make_covarianceMatrix(A))
    vList = store_v(v)
    normalize_vList(vList)

    omegaList = make_omegaList(Phi_list, vList)

    return (xList, xBar, Phi_list, A, lamb, vList, omegaList)

def create_AVPD_data():
    #return xList, xBar,        PhiList, A,           SigmaX,      lamb,        vList, omegaList, k, PhiHat
    return  [],    np.zeros(0), [],      np.zeros(0), np.zeros(0), np.zeros(0), [],     [],       0, np.zeros

def init_AVPD_data(string_list, T):
    xList = []
    for i in string_list:
        xList.append(make_X(i))

    xBar = make_xBar(xList)
    Phi_list = make_PhiList(xList, xBar)
    A = make_A(Phi_list)
    SigmaX = make_covarianceMatrix(A)
    lamb, v = make_Eigen(SigmaX)
    vList = normalize_vList(store_v(v))
    omegaList = make_omegaList(Phi_list, vList)
    k = get_kThresh(lamb, T)
    PhiHat = get_PhiHat(k, omegaList, vList)

    return xList, xBar, Phi_list, A, SigmaX, lamb, vList, omegaList, k, PhiHat

def init_from_xList(xList, T):
    xBar = make_xBar(xList)
    Phi_list = make_PhiList(xList, xBar)
    A = make_A(Phi_list)
    SigmaX = make_covarianceMatrix(A)
    lamb, v = make_Eigen(SigmaX)
    vList = normalize_vList(store_v(v))
    omegaList = make_omegaList(Phi_list, vList)
    k = get_kThresh(lamb, T)
    PhiHat = get_PhiHat(k, omegaList, vList)

    return xBar, Phi_list, A, SigmaX, lamb, vList, omegaList, k, PhiHat

def saveStudentInfo(save_directory, xList, Phi_list, A, SigmaX, lamb, vList, omegaList, PhiHat):
    try:
        os.mkdir(save_directory)
        print("Saving as", save_directory)
    except OSError as error:
        print("Saving as", save_directory)
    with open(save_directory+"/xList.txt", "wb") as fp:
        pickle.dump(xList, fp)
    with open(save_directory+"/Phi_list.txt", "wb") as fp:
        pickle.dump(Phi_list, fp)
    with open(save_directory+"/vList.txt", "wb") as fp:
        pickle.dump(vList, fp)
    with open(save_directory+"/omegaList.txt", "wb") as fp:
        pickle.dump(omegaList, fp)

    np.save(save_directory+"/A.npy", A)
    np.save(save_directory+"/SigmaX.npy", SigmaX)
    np.save(save_directory+"/lambda.npy", lamb)
    np.save(save_directory+"/PhiHat.npy", PhiHat)
    return True

def loadStudentInfo(load_dir, T):
    print("Loading from", load_dir)

    with open(load_dir+"/xList.txt", "rb") as fp:
        xList = pickle.load(fp)
    with open(load_dir+"/Phi_list.txt", "rb") as fp:
        PhiList = pickle.load(fp)
    with open(load_dir+"/vList.txt", "rb") as fp:
        vList = pickle.load(fp)
    with open(load_dir+"/omegaList.txt", "rb") as fp:
        omegaList = pickle.load(fp)

    A = np.load(load_dir+"/A.npy")
    SigmaX = np.load(load_dir+"/SigmaX.npy")
    lamb = np.load(load_dir+"/lambda.npy")
    PhiHat = np.load(load_dir+"/PhiHat.npy")

    xBar = make_xBar(xList)
    k = get_kThresh(lamb, T)

    return xList, xBar, PhiList, A, SigmaX, lamb, vList, omegaList, k, PhiHat

def create_xList_from_dir(dir):
    xList = []
    #data = ''
    with open(dir, 'r') as file:
        data = file.read()
        #print(data)
        #xList = getSentenceList(data)
        xList = makeX_fromParagraph(data)
    #print(xList[0])
    return xList

def integrateEssay(essayList, essayDir, dir):
    essayList.append(essayDir)
    with open(dir, "wb") as fp:
        pickle.dump(essayList, fp)
    return essayList

def loadEssayList(dir):
    essayList = []
    with open(dir, "rb") as fp:
        essayList = pickle.load(fp)
    return essayList

def removeFromListIf(list, sign, value):
    removeList = []
    for i in list:
        if sign == '>' and i < value:
            removeList.append(i)
        elif sign == '<' and i > value:
            removeList.append(i)
        elif sign == '==' and i != value:
            removeList.append(i)
    
    return removeList

def countCriteria(list, sign, value):
    cnt = 0
    for i in list:
        if sign == '>' and i > value:
            cnt+=1
        elif sign == '<' and i < value:
            cnt+=1
        elif sign == '==':
            cnt+=1
    return cnt

def saveList(list, name):
    print("Saving list", name,'\n')
    with open(name, 'wb') as fp:
        pickle.dump(list, fp)
    return True

def loadList(name):
    print("Loading list", name,'\n', flush=True)
    with open(name, 'rb') as fp:
        return pickle.load(fp)

def listContains(list, value):
    for i in list:
        if i == value:
            return True
    return False

def flaggedPercentage(list, threshold):
    return countCriteria(list, '>', threshold)/len(list)